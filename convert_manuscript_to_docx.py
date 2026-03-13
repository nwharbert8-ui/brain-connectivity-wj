"""
Convert brain connectivity manuscript from .txt to .docx (NeuroImage format)
Author: Drake H. Harbert (D.H.H.)
Affiliation: Inner Architecture LLC, Canton, OH
ORCID: 0009-0007-7740-3616
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_FILE = os.path.join(BASE_DIR, "results", "manuscript_draft_brain_connectivity_wj.txt")
OUTPUT_FILE = os.path.join(BASE_DIR, "results", "manuscript_brain_connectivity_wj.docx")

with open(INPUT_FILE, 'r', encoding='utf-8') as f:
    text = f.read()

doc = Document()

# Set margins to 1 inch
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0

# Heading styles
for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Times New Roman'
    heading_style.font.size = Pt(13 if level == 1 else 12)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 0, 0)

# Parse the manuscript text
lines = text.split('\n')
i = 0

# Known section headers that get special treatment
SECTION_HEADERS = {
    'ABSTRACT', 'ACKNOWLEDGMENTS', 'REFERENCES', 'HIGHLIGHTS',
    'FIGURE LEGENDS', 'SUPPLEMENTARY MATERIALS', 'CREDIT AUTHOR STATEMENT'
}

# Track what section we're in for context
current_section = None

while i < len(lines):
    line = lines[i].strip()

    # Skip empty lines
    if not line:
        i += 1
        continue

    # Skip decoration lines
    if line.startswith('=' * 10):
        i += 1
        continue

    # TITLE
    if line == 'TITLE:':
        i += 1
        title_lines = []
        while i < len(lines) and lines[i].strip():
            title_lines.append(lines[i].strip())
            i += 1
        title_text = ' '.join(title_lines)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'

        # Author block
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Drake H. Harbert')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Inner Architecture LLC, Canton, OH')
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('ORCID: 0009-0007-7740-3616')
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Corresponding author: Drake@innerarchitecturellc.com')
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

        doc.add_page_break()
        continue

    # HIGHLIGHTS — special formatting (bullet points)
    if line == 'HIGHLIGHTS':
        doc.add_heading('Highlights', level=1)
        i += 1
        while i < len(lines) and lines[i].strip():
            bullet_line = lines[i].strip()
            if bullet_line.startswith('- '):
                bullet_line = bullet_line[2:]
            p = doc.add_paragraph(bullet_line, style='List Bullet')
            i += 1
        doc.add_paragraph()  # spacing
        continue

    # KEYWORDS — inline paragraph
    if line.startswith('KEYWORDS:'):
        p = doc.add_paragraph()
        run = p.add_run('Keywords: ')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        kw_text = line.replace('KEYWORDS:', '').strip()
        run = p.add_run(kw_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        doc.add_page_break()
        i += 1
        continue

    # DATA AVAILABILITY
    if line == 'DATA AVAILABILITY':
        doc.add_heading('Data Availability', level=1)
        i += 1
        current_section = 'DATA_AVAILABILITY'
        continue

    # CREDIT AUTHOR STATEMENT
    if line == 'CREDIT AUTHOR STATEMENT':
        doc.add_heading('CRediT Author Statement', level=1)
        i += 1
        current_section = 'CREDIT'
        continue

    # Section headings
    if line == 'ABSTRACT':
        doc.add_heading('Abstract', level=1)
        current_section = 'ABSTRACT'
        i += 1
        continue

    if line == 'ACKNOWLEDGMENTS':
        doc.add_heading('Acknowledgments', level=1)
        current_section = 'ACKNOWLEDGMENTS'
        i += 1
        continue

    if line == 'REFERENCES':
        doc.add_page_break()
        doc.add_heading('References', level=1)
        current_section = 'REFERENCES'
        i += 1
        continue

    if line == 'FIGURE LEGENDS':
        doc.add_page_break()
        doc.add_heading('Figure Legends', level=1)
        current_section = 'FIGURE_LEGENDS'
        i += 1
        continue

    if line == 'SUPPLEMENTARY MATERIALS':
        doc.add_page_break()
        doc.add_heading('Supplementary Materials', level=1)
        current_section = 'SUPPLEMENTARY'
        i += 1
        continue

    # Main section headings (1. INTRODUCTION, etc.)
    m = re.match(r'^(\d+)\.\s+([A-Z].+)', line)
    if m and line == line.upper():
        doc.add_heading(f'{m.group(1)}. {m.group(2).title()}', level=1)
        current_section = m.group(2)
        i += 1
        continue

    # Sub-subsection headings (3.6.1, etc.) — check BEFORE subsections
    m = re.match(r'^(\d+\.\d+\.\d+)\s+(.+)', line)
    if m:
        doc.add_heading(f'{m.group(1)} {m.group(2)}', level=3)
        i += 1
        continue

    # Subsection headings (2.1, 3.7, etc.)
    m = re.match(r'^(\d+\.\d+)\s+(.+)', line)
    if m:
        doc.add_heading(f'{m.group(1)} {m.group(2)}', level=2)
        i += 1
        continue

    # Bullet points
    if line.startswith('- '):
        p = doc.add_paragraph(line[2:], style='List Bullet')
        i += 1
        continue

    # Numbered limitation items (in Limitations section)
    m = re.match(r'^(\d+)\.\s+(.+)', line)
    if m and int(m.group(1)) <= 10:
        content = m.group(2)
        while i + 1 < len(lines) and lines[i + 1].startswith('   '):
            i += 1
            content += ' ' + lines[i].strip()
        p = doc.add_paragraph(content, style='List Number')
        i += 1
        continue

    # WJ formula (indented code-like) — render with italic variables and subscripts
    if line.startswith('WJ(A, B)') or line.startswith('(WJ_'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Parse: WJ(A, B) = sum(min(w_A, w_B)) / sum(max(w_A, w_B))
        # Render with italic variables, subscript for A/B after w_
        parts = [
            ('WJ', True, False), ('(', False, False),
            ('A', True, False), (', ', False, False),
            ('B', True, False), (')', False, False),
            (' = \u03A3 min(', False, False),
            ('w', True, False), ('A', False, True),
            (', ', False, False),
            ('w', True, False), ('B', False, True),
            (') / \u03A3 max(', False, False),
            ('w', True, False), ('A', False, True),
            (', ', False, False),
            ('w', True, False), ('B', False, True),
            (')', False, False),
        ]
        for text, is_italic, is_subscript in parts:
            run = p.add_run(text)
            run.italic = is_italic
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if is_subscript:
                run.font.subscript = True
        i += 1
        continue

    # Regular paragraph — collect continuation lines
    if line:
        para_lines = [line]
        while i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if not next_line or next_line.startswith('=') or \
               re.match(r'^\d+\.\s+[A-Z]', next_line) or \
               re.match(r'^\d+\.\d+', next_line) or \
               next_line.startswith('- ') or \
               next_line.startswith('[') or \
               next_line in SECTION_HEADERS or \
               next_line.startswith('KEYWORDS:') or \
               next_line == 'TITLE:' or \
               next_line == 'DATA AVAILABILITY' or \
               next_line == 'CREDIT AUTHOR STATEMENT':
                break
            i += 1
            para_lines.append(lines[i].strip())

        para_text = ' '.join(para_lines)
        para_text = re.sub(r'\s+', ' ', para_text).strip()

        if para_text:
            doc.add_paragraph(para_text)

    i += 1

# Save
doc.save(OUTPUT_FILE)
print(f"Manuscript saved to: {OUTPUT_FILE}")
print(f"File size: {os.path.getsize(OUTPUT_FILE):,} bytes")

# Word count estimate
word_count = len(text.split())
print(f"Approximate word count (full document): {word_count}")
